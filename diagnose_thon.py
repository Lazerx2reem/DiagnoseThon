import os
import re
import logging
from typing import List

import requests
from bs4 import BeautifulSoup
from docx import Document
from dotenv import load_dotenv  # Import dotenv

from utils import openAIPayLoadHelper, parseOpenAIRespone
from constants import SYSTEM_PROMPT, HOUSE_SEASON_8_TITLES, BASE_URL

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables from .env
load_dotenv()

# OpenAI API Key from .env
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

if not OPENAI_API_KEY:
    logger.error("OpenAI API key is missing. Please add it to the .env file.")
    exit(1)


class ParsedObject:
    """
    Reads content from a website page and structures the text read.
    """
    PARSED_MODE = 0
    FULL_TEXT_MODE = 1

    def __init__(self, url: str, episode_name: str) -> None:
        self.url = url
        self.episode_name = episode_name
        self.mode = ParsedObject.PARSED_MODE

        #comment out
        '''
        self.recap_contents = None
        self.zebra_factor = None
        self.zebra_factor_contents = None

        try:
            self.parsed_contents = self.parse_url(self.url)
            self.get_individual_contents()
        except Exception as error:
            logger.error(f"Parsing failed for {url}: {error}")
            logger.info("Using the entire content for LLM prompt generation.")
            self.mode = ParsedObject.FULL_TEXT_MODE
            '''
        self.parsed_contents = self.parse_url(self.url)
        ####
        
    def parse_url(self, url: str) -> str:
        """
        Parses the URL content and returns the text.
        """
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        return soup.get_text()

    def get_individual_contents(self) -> None:
        """
        Extracts the recap contents, zebra factor, and zebra factor contents from the parsed contents.
        """
        self.recap_contents = self.get_contents_between(
            data=self.parsed_contents,
            s1=r"Recap\[\]",
            s2=r"(Clinic Patient\[\]|Clinic Patients\[\])"
        )

        zebra_factor_match = re.search(r"Zebra Factor\s*[:]\s*(\d+)", self.parsed_contents)
        if zebra_factor_match:
            self.zebra_factor = int(zebra_factor_match.group())
        else:
            raise ValueError("Zebra Factor not found.")

        self.zebra_factor_contents = self.get_contents_between(
            data=self.parsed_contents,
            s1=r"Zebra Factor [0-9]+/10\[\]",
            s2=r"(Trivia and cultural references\[\]|Trivia and Cultural References\[\]|Title\[\])"
        )

    @staticmethod
    def get_contents_between(data: str, s1: str, s2: str) -> str:
        """
        Finds and returns the substring between two patterns.
        """
        pattern = f"{s1}[\S\s]*?{s2}"
        matches = re.findall(pattern, data)
        if matches:
            return matches[0]
        else:
            raise ValueError("Contents between the specified patterns not found.")

    def get_prompt(self) -> str:
        """
        Constructs the prompt for the LLM based on the parsed contents.
        """
        '''
        if self.mode == ParsedObject.PARSED_MODE:
            prompt = (
                "Use the information from RECAP and ZEBRA FACTOR for creating your LLM prompt, "
                "required medical answer, and the disease name.\n"
                "### RECAP ###\n"
                f"{self.recap_contents}\n"
                "### ZEBRA FACTOR ###\n"
                f"{self.zebra_factor_contents}\n"
            )
        else:'''
        prompt = (
                "Use details in INFORMATION for creating your LLM prompt, required medical answer, and the disease name.\n"
                "### INFORMATION ###\n"
                f"{self.parsed_contents}\n"
            )
        return prompt

    def save_document(self, generated_text: str, root_dir: str = "./generated/") -> None:
        """
        Saves the generated text to a .docx document.
        """
        document = Document()
        document.add_heading(self.episode_name, 0)
        document.add_paragraph(generated_text)

        # Ensure the directory exists
        os.makedirs(root_dir, exist_ok=True)

        file_name = f"{self.episode_name.replace(' ', '_')}.docx"
        file_path = os.path.join(root_dir, file_name)
        document.save(file_path)
        logger.info(f"Saved to {file_path}")


def main() -> None:
    logger.info("=== Collecting data ===")
    parsed_objs: List[ParsedObject] = []

    # Collect all URLs
    for episode_name in HOUSE_SEASON_8_TITLES:
        url = f"{BASE_URL}{episode_name.replace(' ', '_')}"
        logger.info(f"{episode_name} ::: {url}")
        try:
            parsed_obj = ParsedObject(url, episode_name)
            parsed_objs.append(parsed_obj)
        except Exception as error:
            logger.error(f"Failed to process {url}: {error}")
            logger.info(f"{url} skipped.")

    logger.info("=== Generating medical questions and answers ===")
    # Call the API
    for parsed_obj in parsed_objs:
        prompt = parsed_obj.get_prompt()

        # Format the API payload
        headers, payload = openAIPayLoadHelper(prompt, OPENAI_API_KEY, SYSTEM_PROMPT)

        # Make the POST API call
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
        response.raise_for_status()

        parsed_output = parseOpenAIRespone(response)
        logger.info(parsed_output)
        parsed_obj.save_document(parsed_output)


if __name__ == "__main__":
    main()
